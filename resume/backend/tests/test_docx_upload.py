"""DOCX upload + parse tests."""

from io import BytesIO

import pytest
from docx import Document

from app.services.docx_parser import DocxParser, DocxParseError


def _build_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com | 416-555-0100 | Toronto, ON")
    doc.add_paragraph("Summary")
    doc.add_paragraph(
        "Backend engineer with five years of experience building distributed systems."
    )
    doc.add_paragraph("Experience")
    doc.add_paragraph("Senior Engineer at Acme Corp")
    doc.add_paragraph("2020 - 2024")
    doc.add_paragraph("• Reduced API latency by 40% serving 2M monthly users")
    doc.add_paragraph("• Mentored junior developers across two product teams")
    doc.add_paragraph("Education")
    doc.add_paragraph("BS Computer Science, University of Waterloo, 2019")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, JavaScript, TypeScript, AWS, Docker, Kubernetes")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_parser_extracts_core_fields():
    parser = DocxParser()
    parsed = parser.parse_resume(_build_docx())
    data = parsed.data

    assert data.name == "Jane Doe"
    assert data.email == "jane@example.com"
    assert data.source_format == "docx"

    # Skills are pulled out as a list (subset check — order may vary).
    skills_lower = {s.lower() for s in data.skills}
    assert {"python", "javascript", "aws", "docker", "kubernetes"} <= skills_lower

    # Experience: title/company split, dates parsed.
    assert len(data.experience) >= 1
    exp = data.experience[0]
    assert "Engineer" in exp.title
    assert "Acme" in exp.company
    assert exp.start_date == "2020" and exp.end_date == "2024"

    assert parsed.sections_found  # at least one section detected


def test_docx_parser_rejects_non_docx_bytes():
    with pytest.raises(DocxParseError):
        DocxParser().parse_resume(b"not a docx at all")


def test_docx_parser_rejects_empty_document():
    """A valid .docx with no paragraphs should fail loudly, not return a fake
    'Unknown' resume."""
    doc = Document()
    buf = BytesIO(); doc.save(buf)
    with pytest.raises(DocxParseError):
        DocxParser().parse_resume(buf.getvalue())
