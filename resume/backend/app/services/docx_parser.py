"""DOCX resume parser.

Reads a `.docx` upload via python-docx, flattens paragraphs and table cells
into plain text, and reuses the PDF parser's section detection + per-section
parsers (which work on plain text). This avoids duplicating any of the
experience/education/skills extraction logic.
"""

import logging
from io import BytesIO
from typing import List, Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.models.resume import ResumeData, ParsedResume
from app.services.pdf_parser import get_pdf_parser

logger = logging.getLogger(__name__)


class DocxParseError(Exception):
    """Raised when a .docx file cannot be opened or has no extractable text."""


class DocxParser:
    """Parse a Microsoft Word .docx resume into a ResumeData."""

    def parse_resume(self, docx_bytes: bytes) -> ParsedResume:
        try:
            doc = Document(BytesIO(docx_bytes))
        except PackageNotFoundError as exc:
            raise DocxParseError(
                "Could not open the file as a Word document. It may be corrupt "
                "or not a valid .docx file."
            ) from exc
        except Exception as exc:
            raise DocxParseError(
                "Failed to read the .docx file."
            ) from exc

        full_text = self._extract_text(doc)
        if not full_text.strip():
            raise DocxParseError(
                "No text could be extracted from this document. Please check "
                "that the .docx contains real text (not just images)."
            )

        # Reuse the PDF parser's section detection and per-section parsers —
        # they operate on plain text, so the DOCX flow gets identical structure.
        pdf_parser = get_pdf_parser()
        sections = pdf_parser._detect_sections(full_text)
        resume_data: ResumeData = pdf_parser._parse_sections(sections, full_text)
        resume_data.raw_text = full_text
        resume_data.source_format = "docx"

        return ParsedResume(
            data=resume_data,
            formatting_issues=[],
            confidence=0.75,
            sections_found=list(sections.keys()),
        )

    @staticmethod
    def _extract_text(doc: Document) -> str:
        """Walk paragraphs + table cells; preserve paragraph breaks."""
        chunks: List[str] = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                chunks.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_chunks: List[str] = []
                for cell in row.cells:
                    cell_text = (cell.text or "").strip()
                    if cell_text:
                        row_chunks.append(cell_text)
                if row_chunks:
                    chunks.append("  ".join(row_chunks))
        return "\n".join(chunks)


_docx_parser: Optional[DocxParser] = None


def get_docx_parser() -> DocxParser:
    global _docx_parser
    if _docx_parser is None:
        _docx_parser = DocxParser()
    return _docx_parser
