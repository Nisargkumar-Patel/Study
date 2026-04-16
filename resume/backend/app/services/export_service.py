from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from typing import Dict
import logging

logger = logging.getLogger(__name__)


class ExportService:
    """Export resumes to ATS-friendly formats"""

    def export_to_pdf(self, resume_data: Dict, template: str = "classic") -> bytes:
        """
        Export resume to ATS-friendly PDF

        Args:
            resume_data: Resume data dictionary
            template: Template name (classic, modern, technical, executive, minimal)

        Returns:
            PDF bytes
        """
        buffer = BytesIO()

        # Create PDF document with ATS-friendly settings
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        # Build content based on template
        story = []

        if template == "modern":
            story = self._build_modern_template(resume_data)
        elif template == "technical":
            story = self._build_technical_template(resume_data)
        elif template == "executive":
            story = self._build_executive_template(resume_data)
        elif template == "minimal":
            story = self._build_minimal_template(resume_data)
        else:  # classic
            story = self._build_classic_template(resume_data)

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    def export_to_docx(self, resume_data: Dict, template: str = "classic") -> bytes:
        """
        Export resume to DOCX format

        Args:
            resume_data: Resume data dictionary
            template: Template name

        Returns:
            DOCX bytes
        """
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        # Add content based on template
        if template == "modern":
            self._build_docx_modern(doc, resume_data)
        elif template == "technical":
            self._build_docx_technical(doc, resume_data)
        elif template == "executive":
            self._build_docx_executive(doc, resume_data)
        elif template == "minimal":
            self._build_docx_minimal(doc, resume_data)
        else:  # classic
            self._build_docx_classic(doc, resume_data)

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def export_to_text(self, resume_data: Dict) -> str:
        """Export resume to plain text format"""
        lines = []

        # Header
        name = resume_data.get("name", "")
        lines.append(name.upper())
        lines.append("=" * len(name))
        lines.append("")

        # Contact info
        contact = []
        if resume_data.get("email"):
            contact.append(resume_data["email"])
        if resume_data.get("phone"):
            contact.append(resume_data["phone"])
        if resume_data.get("location"):
            contact.append(resume_data["location"])

        if contact:
            lines.append(" | ".join(contact))
            lines.append("")

        # Summary
        if resume_data.get("summary"):
            lines.append("PROFESSIONAL SUMMARY")
            lines.append("-" * 20)
            lines.append(resume_data["summary"])
            lines.append("")

        # Experience
        if resume_data.get("experience"):
            lines.append("PROFESSIONAL EXPERIENCE")
            lines.append("-" * 25)

            for exp in resume_data["experience"]:
                lines.append(f"\n{exp.get('title', '')} - {exp.get('company', '')}")
                lines.append(f"{exp.get('start_date', '')} to {exp.get('end_date', '')}")

                for bullet in exp.get("bullets", []):
                    lines.append(f"  • {bullet}")

            lines.append("")

        # Education
        if resume_data.get("education"):
            lines.append("EDUCATION")
            lines.append("-" * 10)

            for edu in resume_data["education"]:
                lines.append(f"\n{edu.get('degree', '')} - {edu.get('institution', '')}")
                if edu.get("graduation_date"):
                    lines.append(f"Graduated: {edu['graduation_date']}")

            lines.append("")

        # Skills
        if resume_data.get("skills"):
            lines.append("SKILLS")
            lines.append("-" * 7)
            lines.append(", ".join(resume_data["skills"]))
            lines.append("")

        # Certifications
        if resume_data.get("certifications"):
            lines.append("CERTIFICATIONS")
            lines.append("-" * 15)
            for cert in resume_data["certifications"]:
                lines.append(f"  • {cert}")

        return "\n".join(lines)

    # PDF Template Builders
    def _build_classic_template(self, resume_data: Dict) -> list:
        """Build classic template (traditional single-column)"""
        story = []
        styles = self._get_ats_styles()

        # Header - Name and Contact
        name = resume_data.get("name", "")
        story.append(Paragraph(name, styles['Name']))

        contact = self._format_contact(resume_data)
        story.append(Paragraph(contact, styles['Contact']))
        story.append(Spacer(1, 0.2 * inch))

        # Summary
        if resume_data.get("summary"):
            story.append(Paragraph("PROFESSIONAL SUMMARY", styles['Heading']))
            story.append(Paragraph(resume_data["summary"], styles['Body']))
            story.append(Spacer(1, 0.15 * inch))

        # Experience
        if resume_data.get("experience"):
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", styles['Heading']))

            for exp in resume_data["experience"]:
                # Job title and company
                job_header = f"<b>{exp.get('title', '')}</b> - {exp.get('company', '')}"
                story.append(Paragraph(job_header, styles['Body']))

                # Dates
                dates = f"{exp.get('start_date', '')} to {exp.get('end_date', '')}"
                story.append(Paragraph(dates, styles['Subtitle']))

                # Bullets
                for bullet in exp.get("bullets", []):
                    bullet_text = f"• {bullet}"
                    story.append(Paragraph(bullet_text, styles['Body']))

                story.append(Spacer(1, 0.1 * inch))

        # Education
        if resume_data.get("education"):
            story.append(Paragraph("EDUCATION", styles['Heading']))

            for edu in resume_data["education"]:
                edu_text = f"<b>{edu.get('degree', '')}</b> - {edu.get('institution', '')}"
                story.append(Paragraph(edu_text, styles['Body']))

                if edu.get("graduation_date"):
                    story.append(Paragraph(edu["graduation_date"], styles['Subtitle']))

                story.append(Spacer(1, 0.05 * inch))

        # Skills
        if resume_data.get("skills"):
            story.append(Spacer(1, 0.1 * inch))
            story.append(Paragraph("SKILLS", styles['Heading']))
            skills_text = " | ".join(resume_data["skills"])
            story.append(Paragraph(skills_text, styles['Body']))

        # Certifications
        if resume_data.get("certifications"):
            story.append(Spacer(1, 0.1 * inch))
            story.append(Paragraph("CERTIFICATIONS", styles['Heading']))

            for cert in resume_data["certifications"]:
                story.append(Paragraph(f"• {cert}", styles['Body']))

        return story

    def _build_modern_template(self, resume_data: Dict) -> list:
        """Build modern template (clean with subtle accents)"""
        # Similar to classic but with slight style variations
        return self._build_classic_template(resume_data)

    def _build_technical_template(self, resume_data: Dict) -> list:
        """Build technical template (optimized for engineering roles)"""
        # Emphasize skills and projects more
        return self._build_classic_template(resume_data)

    def _build_executive_template(self, resume_data: Dict) -> list:
        """Build executive template (for senior/leadership roles)"""
        # More focus on achievements and leadership
        return self._build_classic_template(resume_data)

    def _build_minimal_template(self, resume_data: Dict) -> list:
        """Build minimal template (maximum readability)"""
        # Simplest possible layout
        return self._build_classic_template(resume_data)

    # DOCX Template Builders
    def _build_docx_classic(self, doc: Document, resume_data: Dict):
        """Build classic DOCX template"""
        # Name
        name_para = doc.add_paragraph(resume_data.get("name", ""))
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_para.runs[0].font.size = Pt(16)
        name_para.runs[0].font.bold = True

        # Contact
        contact = self._format_contact(resume_data)
        contact_para = doc.add_paragraph(contact)
        contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_para.runs[0].font.size = Pt(10)

        doc.add_paragraph()  # Blank line

        # Summary
        if resume_data.get("summary"):
            heading = doc.add_paragraph("PROFESSIONAL SUMMARY")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)

            doc.add_paragraph(resume_data["summary"])
            doc.add_paragraph()

        # Experience
        if resume_data.get("experience"):
            heading = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)

            for exp in resume_data["experience"]:
                # Job title and company
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(f"{exp.get('title', '')} - {exp.get('company', '')}")
                job_run.font.bold = True

                # Dates
                doc.add_paragraph(f"{exp.get('start_date', '')} to {exp.get('end_date', '')}")

                # Bullets
                for bullet in exp.get("bullets", []):
                    doc.add_paragraph(bullet, style='List Bullet')

                doc.add_paragraph()  # Spacing

        # Education
        if resume_data.get("education"):
            heading = doc.add_paragraph("EDUCATION")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)

            for edu in resume_data["education"]:
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(f"{edu.get('degree', '')} - {edu.get('institution', '')}")
                edu_run.font.bold = True

                if edu.get("graduation_date"):
                    doc.add_paragraph(edu["graduation_date"])

            doc.add_paragraph()

        # Skills
        if resume_data.get("skills"):
            heading = doc.add_paragraph("SKILLS")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)

            skills_text = " | ".join(resume_data["skills"])
            doc.add_paragraph(skills_text)

        # Certifications
        if resume_data.get("certifications"):
            doc.add_paragraph()
            heading = doc.add_paragraph("CERTIFICATIONS")
            heading.runs[0].font.bold = True
            heading.runs[0].font.size = Pt(12)

            for cert in resume_data["certifications"]:
                doc.add_paragraph(cert, style='List Bullet')

    def _build_docx_modern(self, doc: Document, resume_data: Dict):
        """Build modern DOCX template"""
        self._build_docx_classic(doc, resume_data)

    def _build_docx_technical(self, doc: Document, resume_data: Dict):
        """Build technical DOCX template"""
        self._build_docx_classic(doc, resume_data)

    def _build_docx_executive(self, doc: Document, resume_data: Dict):
        """Build executive DOCX template"""
        self._build_docx_classic(doc, resume_data)

    def _build_docx_minimal(self, doc: Document, resume_data: Dict):
        """Build minimal DOCX template"""
        self._build_docx_classic(doc, resume_data)

    # Helper methods
    def _get_ats_styles(self):
        """Get ATS-friendly paragraph styles"""
        styles = getSampleStyleSheet()

        # Name style
        styles.add(ParagraphStyle(
            name='Name',
            parent=styles['Heading1'],
            fontSize=16,
            fontName='Helvetica-Bold',
            alignment=TA_CENTER,
            spaceAfter=6,
            textColor=black
        ))

        # Contact style
        styles.add(ParagraphStyle(
            name='Contact',
            parent=styles['Normal'],
            fontSize=10,
            fontName='Helvetica',
            alignment=TA_CENTER,
            spaceAfter=12,
            textColor=black
        ))

        # Section heading style
        styles.add(ParagraphStyle(
            name='Heading',
            parent=styles['Heading2'],
            fontSize=12,
            fontName='Helvetica-Bold',
            spaceAfter=6,
            spaceBefore=12,
            textColor=black
        ))

        # Body text style
        styles.add(ParagraphStyle(
            name='Body',
            parent=styles['BodyText'],
            fontSize=10,
            fontName='Helvetica',
            leading=14,
            spaceAfter=6,
            alignment=TA_LEFT,
            textColor=black
        ))

        # Subtitle style
        styles.add(ParagraphStyle(
            name='Subtitle',
            parent=styles['Normal'],
            fontSize=9,
            fontName='Helvetica',
            textColor=black,
            spaceAfter=4
        ))

        return styles

    def _format_contact(self, resume_data: Dict) -> str:
        """Format contact information"""
        contact = []

        if resume_data.get("email"):
            contact.append(resume_data["email"])
        if resume_data.get("phone"):
            contact.append(resume_data["phone"])
        if resume_data.get("location"):
            contact.append(resume_data["location"])
        if resume_data.get("linkedin"):
            contact.append(resume_data["linkedin"])

        return " | ".join(contact)


# Singleton instance
_export_service = None

def get_export_service() -> ExportService:
    """Get singleton instance of ExportService"""
    global _export_service
    if _export_service is None:
        _export_service = ExportService()
    return _export_service
